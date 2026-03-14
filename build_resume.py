from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx.oxml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

BLUE = RGBColor(0x1a, 0x0d, 0xab)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x44, 0x44, 0x44)

sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = docx.oxml.OxmlElement('w:r')
    rPr = docx.oxml.OxmlElement('w:rPr')
    c = docx.oxml.OxmlElement('w:color')
    c.set(qn('w:val'), '1a0dab')
    rPr.append(c)
    u = docx.oxml.OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = docx.oxml.OxmlElement('w:sz')
    sz.set(qn('w:val'), str(paragraph.runs[-1].font.size.pt * 2 if paragraph.runs and paragraph.runs[-1].font.size else 22))
    rPr.append(sz)
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = docx.oxml.OxmlElement('w:pBdr')
    bottom = docx.oxml.OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_entry_header(left_text, right_text, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(left_text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    # Right-align date using tab stop
    tab = docx.oxml.OxmlElement('w:tab')
    run2 = p.add_run()
    run2._r.append(tab)
    run3 = p.add_run(right_text)
    run3.font.size = Pt(11)
    run3.font.name = 'Times New Roman'
    # Set right tab stop
    pPr = p._p.get_or_add_pPr()
    tabs = docx.oxml.OxmlElement('w:tabs')
    tab_elem = docx.oxml.OxmlElement('w:tab')
    tab_elem.set(qn('w:val'), 'right')
    tab_elem.set(qn('w:pos'), '9360')  # ~6.5 inches
    tabs.append(tab_elem)
    pPr.append(tabs)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = GRAY

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'

def add_paper_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_paper_detail(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = GRAY
    if italic:
        run.italic = True
    return p

def add_simple_text(text, size=11, bold=False, italic=False, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p

# ===== HEADER =====
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(4)
run = name.add_run('Parker Whitfill')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Times New Roman'

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(0)
run = contact.add_run('Massachusetts Institute of Technology, Department of Economics')
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'

contact2 = doc.add_paragraph()
contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact2.paragraph_format.space_before = Pt(0)
contact2.paragraph_format.space_after = Pt(2)
run = contact2.add_run('50 Memorial Dr, Cambridge, MA 02142')
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'

contact3 = doc.add_paragraph()
contact3.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact3.paragraph_format.space_before = Pt(0)
contact3.paragraph_format.space_after = Pt(4)
run = contact3.add_run('parkerwhitfill@gmail.com  |  whitfill@mit.edu  |  github.com/parkerwhitfill  |  U.S. Citizen')
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'

# ===== SUMMARY =====
add_section_heading('Summary')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
text = (
    'Economics PhD researcher with deep expertise in AI/ML systems, benchmarking, and computational methods. '
    'Published in Nature on large-scale AI evaluation. '
    'Research collaborator with Epoch AI and METR on AI forecasting, compute economics, and algorithmic progress. '
    'Experienced in building ML pipelines for OCR, developing AI benchmarks, and full-stack software development. '
    'Proficient in Python, R, and modern ML frameworks.'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'

# ===== EDUCATION =====
add_section_heading('Education')

add_entry_header('Ph.D. in Economics', '2023 \u2013 Present')
add_subtitle('Massachusetts Institute of Technology, Cambridge, MA')
add_bullet('Research focus: AI capabilities evaluation, benchmark design, compute economics, algorithmic progress')
add_bullet('Coursework in machine learning, econometrics, causal inference, and computational methods')

add_entry_header('B.A. in Economics', '2017 \u2013 2021')
add_subtitle('University of Chicago, Chicago, IL')

# ===== EXPERIENCE =====
add_section_heading('Research & Professional Experience')

add_entry_header('Future of Life Institute Fellow \u2013 Software Engineer', '2025')
add_subtitle('Future of Life Institute, Cambridge, MA')
add_bullet('Designed and built a cooperative technology platform enabling users to coordinate messaging and collective action, from architecture through deployment')
add_bullet('Developed full-stack application using Python and JavaScript, integrating APIs for real-time communication and user coordination workflows')
add_bullet('Implemented backend services, database schema design, and frontend interfaces over a 3-month intensive fellowship focused on AI governance and safety tooling')

add_entry_header('Research Assistant', '2021 \u2013 2023')
add_subtitle('University of Chicago, Chicago, IL')
add_bullet('Designed and implemented machine learning pipelines for optical character recognition (OCR) of historical documents, automating extraction of structured data from large archival corpora')
add_bullet('Built text preprocessing, image segmentation, and classification models using Python (PyTorch, OpenCV, Tesseract) to digitize and parse historical economic records')
add_bullet('Developed data ingestion and cleaning workflows for large-scale historical datasets, enabling quantitative analysis of inequality and institutional change')
add_bullet('Conducted empirical research on economic history, contributing to a publication on the causal effects of WWII on inequality and the social contract in Britain')
add_bullet('Managed end-to-end research pipelines: data collection, feature engineering, statistical modeling, and visualization')

# ===== PUBLICATIONS =====
add_section_heading('Publications & Research')

add_paper_title('A benchmark of expert-level academic questions to assess AI capabilities')
add_paper_detail('Coauthor; consortium paper with >1,000 authors')
add_paper_detail('Nature, 649:1139\u20131146, 2026', italic=True)
add_bullet('Co-developed Humanity\'s Last Exam (HLE), a 3,000-question multi-modal benchmark evaluating frontier AI systems across dozens of academic disciplines')

add_paper_title('Do Benchmarks Overstate AI Capabilities? Evidence from SWE-Bench Verified')
add_paper_detail('Parker Whitfill, Cheryl Wu, Nate Rush, Joel Becker \u2014 Working Paper')
add_bullet('Investigating measurement validity of widely-used software engineering AI benchmarks')

add_paper_title('The Software Intelligence Explosion Debate Needs Experiments')
add_paper_detail('Anson Ho (Epoch AI), Parker Whitfill \u2014 Epoch AI')
add_bullet('Collaboration with Epoch AI analyzing feasibility of recursive AI self-improvement; proposed experimental frameworks for measuring AI R&D productivity')

add_paper_title('Forecasting AI Time Horizon Under Compute Slowdowns')
add_paper_detail('Parker Whitfill, Ben Snodin, Joel Becker (METR) \u2014 arXiv:2511.19492')
add_bullet('Collaboration with METR; built quantitative forecasting models for AI development timelines under hardware constraint scenarios')

add_paper_title('Will Compute Bottlenecks Prevent an Intelligence Explosion?')
add_paper_detail('Parker Whitfill and Cheryl Wu \u2014 arXiv:2507.23181')
add_bullet('Modeled compute supply-demand dynamics for AI scaling; estimated substitutability between compute and algorithmic efficiency')

add_paper_title('Note on Selection Bias in Observational Estimates of Algorithmic Progress')
add_paper_detail('Parker Whitfill \u2014 arXiv:2508.11033')

add_paper_title('Beyond Ordinal Preferences: Why Alignment Needs Cardinal Human Feedback')
add_paper_detail('Parker Whitfill and Stewart Slocum \u2014 arXiv:2508.08486')
add_bullet('Analyzed limitations of RLHF preference models; proposed cardinal feedback mechanisms for AI alignment')

add_paper_title('The Second World War, Inequality and the Social Contract in England')
add_paper_detail('Leander Heldring, James Robinson, Parker Whitfill (equal contribution)')
add_paper_detail('Economica, Volume 89, Issue S1, Pages S137\u2013S159', italic=True)

# ===== TECHNICAL SKILLS =====
add_section_heading('Technical Skills')

skills = [
    ('Languages: ', 'Python, R, SQL, LaTeX, Bash, JavaScript/HTML/CSS'),
    ('ML/AI: ', 'PyTorch, scikit-learn, OpenCV, Tesseract OCR, Hugging Face Transformers, LLM evaluation'),
    ('Data & Compute: ', 'Pandas, NumPy, Stata, large-scale data processing, statistical modeling'),
    ('Tools: ', 'Git, Linux/Unix, cloud computing, Jupyter, VS Code'),
    ('Methods: ', 'Machine learning, deep learning, computer vision (OCR), NLP, causal inference, econometrics, benchmark design & evaluation, time-series forecasting'),
]

for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run = p.add_run(value)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'

# ===== HONORS =====
add_section_heading('Honors & Fellowships')

honors = [
    'Emergent Ventures Grantee (2025)',
    'Stripe Economics of AI Fellowship (2025)',
    'Future of Life Institute Fellowship (2025)',
    '2017 National Debate Champion (TOC LD)',
]
for h in honors:
    add_bullet(h)

doc.save('/Users/parkerwhitfill/parkerwhitfill.github.io/resume_nist.docx')
print('Done: resume_nist.docx')
